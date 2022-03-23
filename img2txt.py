import os
from PIL import Image
from pytesseract import pytesseract
from colorama import Fore
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_LINE_SPACING, WD_BREAK

path_to_tesseract = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
imagesDir = r"images/"
outDir = r'out/'

if not os.path.exists(outDir):
    os.mkdir(outDir)


pytesseract.tesseract_cmd = path_to_tesseract
for index, filename in enumerate(os.listdir(imagesDir)):
    try:
        img = Image.open(imagesDir + filename)
        text = pytesseract.image_to_string(img)
        text = text.strip().split('\n')

        document = Document()
        p = document.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        for section in document.sections:
            section.top_margin = Inches(0.3)
            section.bottom_margin = Inches(0.3)
            section.left_margin = Inches(0.3)
            section.right_margin = Inches(0.3)

        for i in text:
            if i != '':
                s = p.add_run(i)
                s.font.name = 'Times New Roman'
                s.font.size = Pt(13)
                s.add_break(WD_BREAK.TEXT_WRAPPING)

        document.save(outDir+filename.split('.')[0]+'.docx')
        # with open(outDir+filename.split('.')[0]+'.txt', 'w') as f:
        # f.write(text)
        print(f"[{index}]\t" + Fore.LIGHTGREEN_EX +
              filename + ": SUCCESS" + Fore.RESET)
    except:
        print(f"[{index}]\t" + Fore.LIGHTRED_EX +
              filename + ": ERROR" + Fore.RESET)
